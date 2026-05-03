# app/routers/gst_authorized_signatory.py
# -*- coding: utf-8 -*-
from fastapi import APIRouter, HTTPException, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any, List
import json, urllib.parse, re

from docx import Document
try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
except Exception:
    HAS_DOCXTPL = False

from ..settings import COMPANIES_ROOT

router = APIRouter(prefix="/api/gst/authorized-signatory", tags=["gst-authorized-signatory"])

TEMPLATE_DOCX = Path(r"E:\Bots\docgen_backend_minimal_v2\templates\word\GST_Registration_Board_Resolution.docx")
if not TEMPLATE_DOCX.exists():
    raise RuntimeError(f"Template not found at {TEMPLATE_DOCX}")

def _log(msg: str): print(f"[gst_authorized_signatory] {msg}")

# ---------- path + master helpers ----------
_slug_rx = re.compile(r"[^a-z0-9]+")
def _slugify(s: str) -> str:
    return _slug_rx.sub("-", (s or "").strip().lower()).strip("-")

def _find_company_dir(company_name: str) -> Path:
    exact = COMPANIES_ROOT / company_name
    if exact.exists():
        return exact
    wanted_ci = (company_name or "").casefold()
    for d in COMPANIES_ROOT.iterdir():
        if d.is_dir() and d.name.casefold() == wanted_ci:
            return d
    want_slug = _slugify(company_name)
    for d in COMPANIES_ROOT.iterdir():
        if d.is_dir() and _slugify(d.name) == want_slug:
            return d
    raise HTTPException(status_code=404, detail="Company folder not found")

def _safe_master_path(company_name: str) -> Path:
    cdir = _find_company_dir(company_name)
    masters = sorted(cdir.glob("_master_*.json"))
    if not masters:
        raise HTTPException(status_code=404, detail="Master not found for company")
    return masters[-1]

def _load_master(company_name: str) -> Dict[str, Any]:
    mp = _safe_master_path(company_name)
    _log(f"using master: {mp}")
    try:
        return json.loads(mp.read_text(encoding="utf-8"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read master: {e}")

def _save_master(company_name: str, data: Dict[str, Any]) -> None:
    mp = _safe_master_path(company_name)
    mp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    _log(f"master updated: {mp}")

def _string_address(addr: Any) -> str:
    if not addr: return "-"
    if isinstance(addr, str): return addr
    if isinstance(addr, dict):
        parts = [addr.get("line1"), addr.get("line2"), addr.get("city"), addr.get("state"), addr.get("pincode")]
        return ", ".join([str(x) for x in parts if x])
    return str(addr)

def _company_ctx_full(master: Dict[str, Any], company_name: str) -> Dict[str, Any]:
    base = master.get("company") or {}
    reg_addr = master.get("registered_address") or master.get("address") \
               or base.get("registered_address") or base.get("address")
    pan_obj = master.get("pan") or base.get("pan") or {}
    return {
        "name": company_name,
        "cin": master.get("cin") or base.get("cin") or "-",
        "tan": master.get("tan") or base.get("tan") or "-",
        "email": master.get("email") or base.get("email") or "-",
        "registered_address": _string_address(reg_addr) or "-",
        "pan": pan_obj if isinstance(pan_obj, dict) else {"pan_number": pan_obj or "-"},
    }

def _fmt_date(d: Optional[str]) -> str:
    if not d: return datetime.now().strftime("%d %b %Y")
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%Y/%m/%d"):
        try: return datetime.strptime(d.strip(), fmt).strftime("%d %b %Y")
        except Exception: pass
    return d

def _ensure_out_dir(company_name: str) -> Path:
    cdir = _find_company_dir(company_name)     # save in company root
    cdir.mkdir(parents=True, exist_ok=True)
    return cdir

def _public_download_url(company_name: str, out_path: Path) -> str:
    qc = urllib.parse.quote(company_name, safe="")
    qf = urllib.parse.quote(out_path.name, safe="")
    return f"/api/gst/authorized-signatory/download?company_name={qc}&file={qf}"

# ---------- replacers ----------
def _replace_in_paragraphs(paragraphs, fn):
    for p in paragraphs:
        if not p.text: continue
        new = fn(p.text)
        if new != p.text:
            p.text = new

def _walk_tables(tables, fn):
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, fn)
                _walk_tables(cell.tables, fn)
# catch {{ extra.date or today }} / {{extra.date or today}} / {{ extra.date }} / {{ today }}
RX_EXTRA_OR_TODAY = re.compile(r"\{\{\s*extra\.date\s*(?:\|\|\s*|(?:\s+or\s+))?today\s*\}\}", re.IGNORECASE)
RX_EXTRA_DATE_ONLY = re.compile(r"\{\{\s*extra\.date\s*\}\}", re.IGNORECASE)
RX_TODAY_ONLY = re.compile(r"\{\{\s*today\s*\}\}", re.IGNORECASE)
RX_COMPANY_NAME   = re.compile(r"\{\{\s*company\.name\s*\}\}")
RX_COMPANY_CIN    = re.compile(r"\{\{\s*company\.cin\s*or\s*'-'\s*\}\}")
RX_COMPANY_TAN    = re.compile(r"\{\{\s*company\.tan\s*or\s*'-'\s*\}\}")
RX_COMPANY_EMAIL  = re.compile(r"\{\{\s*company\.email\s*or\s*'-'\s*\}\}")
RX_COMPANY_ADDR   = re.compile(r"\{\{\s*company\.registered_address\s*or\s*'-'\s*\}\}?")
RX_PAN_COMPLEX    = re.compile(
    r"\{\{\s*company\.pan\.pan_number\s*if\s*company\.pan\s*else\s*\(pan\.pan_number\s*if\s*pan\s*else\s*'-'\)\s*\}\}"
)
RX_BLOCK = re.compile(r"\{\%.*?\%\}")  # remove stray {% ... %}

def _mk_replacer(vals: Dict[str, str]):
    direct_map = {
        "[DATE]": vals["date"],
        "[Place]": vals["place"],
        "[PLACE]": vals["place"],
        "[Director Name]": vals["signatory_name"],
        "[Name of Director/Authorized Signatory]": vals["signatory_name"],
        "[AUTHORIZED SIGNATORY NAME]": vals["signatory_name"],
        "[Authorized Person Name]": vals["signatory_name"],
        "[DIN]": vals["signatory_din"],
        "[DESIGNATION]": vals["signatory_desig"],
        "[COMPANY NAME]": vals["company_name"],
        "[Company Name]": vals["company_name"],
        "[ADDRESS]": vals["company_addr"],
        "[CIN]": vals["company_cin"],
        "[For Company Name]": vals["company_name"],
        "{% for pair in directors_pairs %}": "",
        "{% endfor %}": "",
        "06 Nov 2025": vals["date"],  # overwrite any sample date lingering in template
    }
    def apply_all(text: str) -> str:
        # simple token replacements
        for k, v in direct_map.items():
            if k in text:
                text = text.replace(k, v)

        # jinja-like {{ ... }} date variants
        text = RX_EXTRA_OR_TODAY.sub(vals["date"], text)   # {{ extra.date or today }} or {{ extra.date || today }}
        text = RX_EXTRA_DATE_ONLY.sub(vals["date"], text)  # {{ extra.date }}
        text = RX_TODAY_ONLY.sub(vals["date"], text)       # {{ today }}

        # other jinja-like singletons we already handle
        text = RX_COMPANY_NAME.sub(vals["company_name"], text)
        text = RX_COMPANY_CIN.sub(vals["company_cin"] or "-", text)
        text = RX_COMPANY_TAN.sub(vals["company_tan"] or "-", text)
        text = RX_COMPANY_EMAIL.sub(vals["company_email"] or "-", text)
        text = RX_COMPANY_ADDR.sub(vals["company_addr"] or "-", text)
        text = RX_PAN_COMPLEX.sub(vals["company_pan"] or "-", text)

        # remove any stray {% ... %} blocks
        text = RX_BLOCK.sub("", text)
        return text
    return apply_all


# ---------- directors loop fallback ----------
def _norm_dir(d: Dict[str, Any]) -> Dict[str, str]:
    if not d: return {}
    return {
        "name": d.get("name") or d.get("full_name") or d.get("director_name") or "",
        "din":  d.get("din") or d.get("DIN") or d.get("din_number") or "",
    }

def _render_directors_cells(doc: Document, directors: List[Dict[str, Any]]):
    """
    Fallback when docxtpl didn't execute the {% for pair in directors_pairs %} loop.
    Works even if {% for %} and {% endfor %} are not in the same cell.
    """
    if not directors: return
    pairs = []
    for i in range(0, len(directors), 2):
        left = _norm_dir(directors[i]) if i < len(directors) else {}
        right = _norm_dir(directors[i+1]) if i+1 < len(directors) else {}
        pairs.append((left, right))

    for t in doc.tables:
        # Try to detect a two-column "Director 1 / Director 2" area by token presence
        for row in t.rows:
            cells = list(row.cells)
            for ci, cell in enumerate(cells):
                cell_text = "\n".join(p.text or "" for p in cell.paragraphs)
                is_left = ("pair[0]" in cell_text)
                is_right = ("pair[1]" in cell_text)
                if not (is_left or is_right):
                    continue

                cell.text = ""
                for idx, (left, right) in enumerate(pairs):
                    person = left if is_left else right
                    name = person.get("name", "") or ""
                    din  = person.get("din", "") or ""
                    cell.add_paragraph(name)
                    cell.add_paragraph(f"DIN: {din}")
                    cell.add_paragraph("____________________")
                    if idx != len(pairs) - 1:
                        cell.add_paragraph("")

# ---------- models ----------
class Signatory(BaseModel):
    name: str = Field(..., description="Selected director full name")
    din: Optional[str] = None
    designation: Optional[str] = "Director"

class Payload(BaseModel):
    company_name: str
    meeting_date: Optional[str] = None
    meeting_place: Optional[str] = None
    signatory: Signatory

# ---------- lightweight reads for the frontend ----------
@router.get("/master")
def get_master(company_name: str = Query(...)):
    return _load_master(company_name)

@router.get("/directors")
def get_directors(company_name: str = Query(...)):
    data = _load_master(company_name)
    directors = data.get("directors") or (data.get("company") or {}).get("directors") or []
    return {"company_name": company_name, "directors": directors}

# ---------- preview ----------
@router.post("/preview")
def preview_resolution(body: Payload):
    master = _load_master(body.company_name)
    company = _company_ctx_full(master, body.company_name)
    html = f"""
    <div style="padding:20px 24px; font-family:Segoe UI, Roboto, Arial, sans-serif; color:#111;">
      <h2 style="margin:0 0 8px 0;">Certified True Copy of the Resolution</h2>
      <div style="color:#555; font-size:13px; margin-bottom:12px;">
        of the Board of Directors of <strong>{company['name']}</strong><br/>
        CIN: {company['cin']}<br/>
        Regd. Office: {company['registered_address']}
      </div>
      <div style="margin:12px 0 16px 0; line-height:1.6; font-size:14px;">
        <div><strong>Meeting Date:</strong> {_fmt_date(body.meeting_date)}</div>
        <div><strong>Meeting Place:</strong> {body.meeting_place or "Registered Office"}</div>
      </div>
      <p style="line-height:1.7; font-size:14px; margin:12px 0;">
        “RESOLVED THAT Mr./Ms. <strong>{body.signatory.name}</strong> (DIN: {body.signatory.din or '-'}) be and is
        hereby appointed as the <strong>Authorised Signatory</strong> for GST matters of the Company…”
      </p>
    </div>
    """
    return {"html": html}

# ---------- generate (fill) ----------
@router.post("/fill")
def fill_resolution(body: Payload):
    master = _load_master(body.company_name)
    company = _company_ctx_full(master, body.company_name)

    directors = master.get("directors") or []
    ctx = {
        "company": company,
        "pan": company["pan"],
        "today": datetime.now().strftime("%d %b %Y"),
        "extra": {"date": _fmt_date(body.meeting_date)},
        "meeting": {"place": body.meeting_place or "Registered Office"},
        "directors_pairs": [[directors[i], directors[i+1] if i+1 < len(directors) else None]
                            for i in range(0, len(directors), 2)],
        "signatory": body.signatory.dict(),
    }

    out_dir = _ensure_out_dir(body.company_name)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    temp_docx = out_dir / f"~tmp_gst_signatory_{ts}.docx"
    final_docx = out_dir / f"bod_gst_authorized_signatory_{ts}.docx"

    # 1) Try Jinja
    jinja_ok = False
    if HAS_DOCXTPL:
        try:
            tpl = DocxTemplate(str(TEMPLATE_DOCX))
            tpl.render(ctx)
            tpl.save(str(temp_docx))
            jinja_ok = True
            _log(f"docxtpl rendered → {temp_docx}")
        except Exception as e:
            _log(f"docxtpl render error: {e!r}; using original template for fallback")
            temp_docx = TEMPLATE_DOCX
    else:
        temp_docx = TEMPLATE_DOCX

    # 2) Build value map and apply replacements
    vals = {
        "date": _fmt_date(body.meeting_date),
        "place": body.meeting_place or "Registered Office",
        "signatory_name": body.signatory.name,
        "signatory_din": body.signatory.din or "-",
        "signatory_desig": body.signatory.designation or "Director",
        "company_name": company["name"],
        "company_addr": company["registered_address"],
        "company_cin": company["cin"] or "-",
        "company_tan": company["tan"] or "-",
        "company_email": company["email"] or "-",
        "company_pan": (company.get("pan") or {}).get("pan_number") or "-",
    }
    replacer = _mk_replacer(vals)

    doc = Document(str(temp_docx))
    if not jinja_ok:
        _render_directors_cells(doc, directors)
    _replace_in_paragraphs(doc.paragraphs, replacer)
    _walk_tables(doc.tables, replacer)
    for section in doc.sections:
        _replace_in_paragraphs(section.header.paragraphs, replacer)
        _walk_tables(section.header.tables, replacer)
        _replace_in_paragraphs(section.footer.paragraphs, replacer)
        _walk_tables(section.footer.tables, replacer)

    doc.save(str(final_docx))
    _log(f"final DOCX → {final_docx}")

    # 3) Update master
    ts_now = datetime.now().strftime("%Y%m%d_%H%M%S")
    rec = {
        "name": body.signatory.name,
        "din": body.signatory.din,
        "designation": body.signatory.designation or "Director",
        "meeting_date": vals["date"],
        "meeting_place": vals["place"],
        "docx": str(final_docx),
        "ts": ts_now,
        "note": "authorised signatory appointed"
    }
    master["authorized_signatory"] = rec
    master["authorised_signatory"] = rec
    meta = master.get("meta") or {}
    meta["last_updated"] = ts_now
    master["meta"] = meta
    plog = master.get("process_log") or []
    plog.append({
        "process": "GST_AUTH_SIGNATORY_APPOINTED",
        "ts": ts_now,
        "inputs": {
            "meeting_date": vals["date"],
            "meeting_place": vals["place"],
            "signatory": {"name": rec["name"], "din": rec["din"], "designation": rec["designation"]}
        },
        "outputs": {"docx": str(final_docx)},
        "status": "SUCCESS"
    })
    master["process_log"] = plog
    _save_master(body.company_name, master)

    return {
        "output": str(final_docx),
        "download_url": _public_download_url(body.company_name, final_docx),
        "company": company["name"],
        "updated_master": True,
        "jinja_ok": jinja_ok
    }

@router.post("/generate")
def generate_resolution(body: Payload):
    return fill_resolution(body)

@router.get("/download")
def download_filled(company_name: str = Query(...), file: str = Query(...)):
    company_dir = _find_company_dir(company_name).resolve()
    candidates = [
        (company_dir / file).resolve(),
        (company_dir / "outputs" / "gst" / file).resolve(),  # legacy
    ]
    target = None
    for p in candidates:
        if str(p).startswith(str(company_dir)) and p.exists():
            target = p; break
    if not target:
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(
        path=str(target),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=target.name
    )

# ---------- Compatibility router expected by main.py ----------
compat = APIRouter()

@compat.get("/api/company/master")
def _compat_master(company_name: str = Query(...)):
    return get_master(company_name)

@compat.get("/api/company/directors")
def _compat_directors(company_name: str = Query(...)):
    return get_directors(company_name)
